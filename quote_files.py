# quote_files.py
# -------------------------------------------------------------------
# Renderizadores EDITABLES para las cotizaciones de JRS:
#   - generar_docx_cotizacion(quote_data, output_path, logo_path=None) -> str
#   - generar_xlsx_cotizacion(quote_data, output_path) -> str
#
# Consumen el MISMO diccionario `quote_data` que quote_pdf.py, de modo
# que son "drop-in" paralelos al PDF. Mismo esquema, misma fuente de
# verdad. Asi, PDF / DOCX / XLSX siempre cuentan la misma historia.
#
# Filosofia compartida con el PDF: estos renderizadores NO hacen
# matematicas. Los importes llegan ya formateados como strings
# ("$21,000", "$18,500 - $26,500"). El numero que ve el cliente es
# exactamente el que calculo el agente.
#
# DECISION DE CONFIDENCIALIDAD (Seccion 18 del system prompt):
# El XLSX es CLIENT-FACING y limpio. NO incluye el metodo crew-day /
# unit-price, ni margenes, ni nivel de confianza interno. Ese archivo
# puede terminar reenviado a un GC; la calculadora interna NUNCA viaja
# dentro de un archivo client-facing. Si en el futuro Richard quiere un
# "backup interno" con los dos metodos y la confianza, sera un archivo
# INTERNO aparte, no este.
#
# CLAVES DE quote_data (todas opcionales salvo las del encabezado;
# las secciones vacias simplemente no se dibujan). Identico a quote_pdf,
# con dos claves OPCIONALES nuevas para soportar cotizaciones con
# material sin mentir en las etiquetas:
#   "basis":              ej. "Labor Only" | "Labor + JRS-Furnished
#                         Materials"  (subtitulo; default "Labor-Only")
#   "materials_subtotal": ej. "$13,900"   (si existe, se dibuja su fila)
# -------------------------------------------------------------------

# ===================================================================
# Paleta de marca JRS (igual que quote_pdf.py)
# ===================================================================
NAVY_HEX = "1F3A5F"
GOLD_HEX = "C8922A"
LIGHT_HEX = "EEF2F7"
DARK_HEX = "222222"
MUTED_HEX = "666666"


# ===================================================================
# Utilidades comunes
# ===================================================================
def _basis_label(quote_data: dict) -> str:
    """Subtitulo honesto. Si no se especifica basis, asumimos labor-only."""
    return str(quote_data.get("basis") or "Labor-Only Subcontractor Quote")


def _lineas_summary(quote_data: dict):
    """El project_summary puede traer varios parrafos separados por \\n."""
    texto = str(quote_data.get("project_summary", "") or "")
    return [p.strip() for p in texto.split("\n") if p.strip()]


# ===================================================================
# 1) DOCX  -- python-docx
# ===================================================================
def generar_docx_cotizacion(quote_data: dict, output_path: str,
                            logo_path: str = None) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY = RGBColor(0x1F, 0x3A, 0x5F)
    GOLD = RGBColor(0xC8, 0x92, 0x2A)
    DARK = RGBColor(0x22, 0x22, 0x22)
    MUTED = RGBColor(0x66, 0x66, 0x66)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    # ---- helpers internos ----
    def _shade(elem, hex_color):
        """Pinta el fondo de una celda o parrafo (w:shd)."""
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        elem.append(shd)

    def _set_cell_bg(cell, hex_color):
        _shade(cell._tc.get_or_add_tcPr(), hex_color)

    def _no_space(p):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)

    def _seccion(titulo):
        """Barra navy con texto blanco, ancho completo (como el PDF)."""
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = tbl.rows[0].cells[0]
        _set_cell_bg(cell, NAVY_HEX)
        p = cell.paragraphs[0]
        _no_space(p)
        run = p.add_run(titulo)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = WHITE
        # margen superior antes de cada seccion
        cell.paragraphs[0].paragraph_format.space_before = Pt(8)
        return tbl

    def _bullets(items):
        for it in items:
            p = doc.add_paragraph(style="List Bullet")
            _no_space(p)
            r = p.add_run(str(it))
            r.font.size = Pt(9.5)
            r.font.color.rgb = DARK

    # ---- documento ----
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Encabezado: logo/empresa + meta de cotizacion
    head = doc.add_table(rows=1, cols=2)
    head.columns[0].width = Inches(4.4)
    head.columns[1].width = Inches(2.4)
    izq = head.rows[0].cells[0]
    p = izq.paragraphs[0]
    _no_space(p)
    if logo_path:
        try:
            p.add_run().add_picture(logo_path, width=Inches(1.4))
        except Exception:
            pass
    rname = izq.add_paragraph().add_run("JRS RETAIL SERVICES")
    rname.bold = True
    rname.font.size = Pt(18)
    rname.font.color.rgb = NAVY
    rtag = izq.add_paragraph().add_run(
        "Commercial Retail Construction / Retail Remodel Subcontractor")
    rtag.font.size = Pt(8.5)
    rtag.font.color.rgb = MUTED

    der = head.rows[0].cells[1]
    pm = der.paragraphs[0]
    pm.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _no_space(pm)
    rq = pm.add_run("QUOTE\n")
    rq.bold = True
    rq.font.size = Pt(11)
    rq.font.color.rgb = NAVY
    rmeta = pm.add_run(
        f"# {quote_data.get('quote_number', 'TBD')}\n"
        f"{quote_data.get('date', '')}")
    rmeta.font.size = Pt(9)
    rmeta.font.color.rgb = DARK

    # Linea dorada divisoria
    div = doc.add_paragraph()
    _no_space(div)
    pPr = div._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), GOLD_HEX)
    pbdr.append(bottom)
    pPr.append(pbdr)

    # Bloque meta (Prepared for / Project / etc.)
    meta_pairs = [
        ("PREPARED FOR", quote_data.get("prepared_for", "")),
        ("PROJECT", quote_data.get("project_name", "")),
        ("STORE #", quote_data.get("store_number", "")),
        ("LOCATION", quote_data.get("location", "")),
        ("PREPARED BY", quote_data.get("prepared_by",
                                       "Richard Bodington / JRS Retail Services")),
        ("PHONE", quote_data.get("phone", "832-361-6551")),
        ("BASIS", _basis_label(quote_data)),
    ]
    mt = doc.add_table(rows=0, cols=2)
    mt.columns[0].width = Inches(1.6)
    mt.columns[1].width = Inches(5.2)
    for label, val in meta_pairs:
        if not str(val).strip():
            continue
        row = mt.add_row().cells
        rp = row[0].paragraphs[0]
        _no_space(rp)
        rl = rp.add_run(label)
        rl.bold = True
        rl.font.size = Pt(9)
        rl.font.color.rgb = NAVY
        vp = row[1].paragraphs[0]
        _no_space(vp)
        rv = vp.add_run(str(val))
        rv.font.size = Pt(9)
        rv.font.color.rgb = DARK

    # 1. PROJECT SUMMARY
    summ = _lineas_summary(quote_data)
    if summ:
        _seccion("1. PROJECT SUMMARY")
        for parrafo in summ:
            pp = doc.add_paragraph()
            _no_space(pp)
            rr = pp.add_run(parrafo)
            rr.font.size = Pt(9.5)
            rr.font.color.rgb = DARK

    # 2. SCOPE OF WORK INCLUDED
    if quote_data.get("scope_items"):
        _seccion("2. SCOPE OF WORK INCLUDED")
        _bullets(quote_data["scope_items"])

    # 3. ESTIMATE BREAKDOWN
    if quote_data.get("line_items"):
        _seccion("3. ESTIMATE BREAKDOWN")
        cols = ["#", "DESCRIPTION", "QTY", "UNIT", "UNIT PRICE", "LINE TOTAL"]
        widths = [Inches(0.35), Inches(3.0), Inches(0.7),
                  Inches(0.6), Inches(1.0), Inches(1.1)]
        tbl = doc.add_table(rows=1, cols=len(cols))
        tbl.style = "Table Grid"
        for i, (c, w) in enumerate(zip(cols, widths)):
            cell = tbl.rows[0].cells[i]
            cell.width = w
            _set_cell_bg(cell, NAVY_HEX)
            cp = cell.paragraphs[0]
            _no_space(cp)
            if i >= 4:
                cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cr = cp.add_run(c)
            cr.bold = True
            cr.font.size = Pt(8.5)
            cr.font.color.rgb = WHITE
        for idx, li in enumerate(quote_data["line_items"], start=1):
            vals = [
                str(idx),
                str(li.get("description", "")),
                str(li.get("qty", "")),
                str(li.get("unit", "")),
                str(li.get("unit_price", "")),
                str(li.get("line_total", "")),
            ]
            cells = tbl.add_row().cells
            for i, (v, w) in enumerate(zip(vals, widths)):
                cells[i].width = w
                cp = cells[i].paragraphs[0]
                _no_space(cp)
                if i >= 4 or i == 2:
                    cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cr = cp.add_run(v)
                cr.font.size = Pt(8.5)
                cr.font.color.rgb = DARK

    # 4. TRAVEL / LODGING / EQUIPMENT
    if quote_data.get("travel_items"):
        _seccion("4. TRAVEL / LODGING / EQUIPMENT (ALLOWANCES)")
        tt = doc.add_table(rows=0, cols=2)
        tt.columns[0].width = Inches(5.7)
        tt.columns[1].width = Inches(1.1)
        for t in quote_data["travel_items"]:
            cells = tt.add_row().cells
            a = cells[0].paragraphs[0]; _no_space(a)
            ra = a.add_run(str(t.get("description", "")))
            ra.font.size = Pt(9); ra.font.color.rgb = DARK
            b = cells[1].paragraphs[0]; _no_space(b)
            b.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rb = b.add_run(str(t.get("amount", "")))
            rb.font.size = Pt(9); rb.font.color.rgb = DARK

    # 5. TOTAL PRICE
    _seccion("5. TOTAL PRICE")
    tot = doc.add_table(rows=0, cols=2)
    tot.columns[0].width = Inches(5.0)
    tot.columns[1].width = Inches(1.8)

    def _fila_total(label, valor, bold=False, navy=False):
        cells = tot.add_row().cells
        a = cells[0].paragraphs[0]; _no_space(a)
        ra = a.add_run(label)
        ra.bold = bold; ra.font.size = Pt(10 if bold else 9.5)
        ra.font.color.rgb = NAVY if navy else DARK
        b = cells[1].paragraphs[0]; _no_space(b)
        b.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rb = b.add_run(str(valor))
        rb.bold = bold; rb.font.size = Pt(12 if bold else 9.5)
        rb.font.color.rgb = NAVY if navy else DARK

    if quote_data.get("labor_subtotal"):
        _fila_total("Labor subtotal", quote_data["labor_subtotal"])
    if quote_data.get("materials_subtotal"):
        _fila_total("Materials subtotal (JRS-furnished)",
                    quote_data["materials_subtotal"])
    if quote_data.get("travel_subtotal"):
        _fila_total("Travel / lodging / equipment", quote_data["travel_subtotal"])
    _fila_total("TOTAL", quote_data.get("total_text", "TBD"),
                bold=True, navy=True)

    # 6-9 listas
    if quote_data.get("assumptions"):
        _seccion("6. ASSUMPTIONS"); _bullets(quote_data["assumptions"])
    if quote_data.get("exclusions"):
        _seccion("7. EXCLUSIONS"); _bullets(quote_data["exclusions"])
    if quote_data.get("clarifications"):
        _seccion("8. CLARIFICATIONS / OPEN ITEMS")
        _bullets(quote_data["clarifications"])
    if quote_data.get("terms"):
        _seccion("9. TERMS"); _bullets(quote_data["terms"])

    # 10. AUTHORIZATION
    _seccion("10. AUTHORIZATION")
    for label in ("Accepted by:", "Name / Title:", "Date:"):
        p = doc.add_paragraph()
        _no_space(p)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"{label}  _________________________")
        r.font.size = Pt(10); r.font.color.rgb = DARK

    # Nota de compliance (opcional)
    if quote_data.get("compliance_note"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        r = p.add_run("COMPLIANCE NOTE: " + str(quote_data["compliance_note"]))
        r.italic = True; r.font.size = Pt(8); r.font.color.rgb = MUTED

    # Pie de pagina (neutral, sin sobre-afirmar "labor-only")
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        "JRS Retail Services  |  Commercial Retail Construction Quote  |  Confidential")
    fr.font.size = Pt(7.5); fr.font.color.rgb = MUTED

    doc.save(output_path)
    return output_path


# ===================================================================
# 2) XLSX  -- openpyxl  (client-facing, limpio, editable)
# ===================================================================
def generar_xlsx_cotizacion(quote_data: dict, output_path: str) -> str:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Quote"

    navy_fill = PatternFill("solid", fgColor=NAVY_HEX)
    light_fill = PatternFill("solid", fgColor=LIGHT_HEX)
    white_bold = Font(bold=True, color="FFFFFF", size=10)
    navy_bold = Font(bold=True, color=NAVY_HEX, size=11)
    dark = Font(color=DARK_HEX, size=10)
    muted = Font(color=MUTED_HEX, size=8, italic=True)
    thin = Side(style="thin", color="CCD4DE")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    wrap = Alignment(wrap_text=True, vertical="top")
    right = Alignment(horizontal="right", vertical="top")
    left = Alignment(horizontal="left", vertical="top")

    NCOLS = 6  # A..F
    widths = [5, 46, 9, 8, 13, 15]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    r = 1

    def _section(text):
        nonlocal r
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=NCOLS)
        for c in range(1, NCOLS + 1):
            cell = ws.cell(row=r, column=c)
            cell.fill = navy_fill
        cell = ws.cell(row=r, column=1, value=text)
        cell.font = white_bold
        cell.alignment = left
        r += 1

    def _meta(label, value):
        nonlocal r
        if not str(value).strip():
            return
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
        lc = ws.cell(row=r, column=1, value=label)
        lc.font = Font(bold=True, color=NAVY_HEX, size=9)
        lc.alignment = left
        ws.merge_cells(start_row=r, start_column=3, end_row=r, end_column=NCOLS)
        v = ws.cell(row=r, column=3, value=str(value))
        v.font = dark
        v.alignment = left
        r += 1

    def _wrapline(text):
        nonlocal r
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=NCOLS)
        cell = ws.cell(row=r, column=1, value=str(text))
        cell.font = dark
        cell.alignment = wrap
        r += 1

    # ---- Encabezado de marca ----
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=NCOLS)
    h = ws.cell(row=r, column=1, value="JRS RETAIL SERVICES")
    h.font = Font(bold=True, color=NAVY_HEX, size=16)
    r += 1
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=NCOLS)
    t = ws.cell(row=r, column=1,
                value="Commercial Retail Construction / Retail Remodel Subcontractor")
    t.font = Font(color=MUTED_HEX, size=9)
    r += 2

    # ---- Meta ----
    _meta("Quote #", quote_data.get("quote_number", "TBD"))
    _meta("Date", quote_data.get("date", ""))
    _meta("Prepared For", quote_data.get("prepared_for", ""))
    _meta("Project", quote_data.get("project_name", ""))
    _meta("Store #", quote_data.get("store_number", ""))
    _meta("Location", quote_data.get("location", ""))
    _meta("Prepared By", quote_data.get("prepared_by",
                                        "Richard Bodington / JRS Retail Services"))
    _meta("Phone", quote_data.get("phone", "832-361-6551"))
    _meta("Basis", _basis_label(quote_data))
    r += 1

    # ---- 1. Project summary ----
    summ = _lineas_summary(quote_data)
    if summ:
        _section("1. PROJECT SUMMARY")
        for parrafo in summ:
            _wrapline(parrafo)
        r += 1

    # ---- 2. Scope ----
    if quote_data.get("scope_items"):
        _section("2. SCOPE OF WORK INCLUDED")
        for it in quote_data["scope_items"]:
            _wrapline("•  " + str(it))
        r += 1

    # ---- 3. Estimate breakdown ----
    if quote_data.get("line_items"):
        _section("3. ESTIMATE BREAKDOWN")
        headers = ["#", "Description", "Qty", "Unit", "Unit Price", "Line Total"]
        for c, htext in enumerate(headers, start=1):
            cell = ws.cell(row=r, column=c, value=htext)
            cell.fill = navy_fill
            cell.font = white_bold
            cell.border = border
            cell.alignment = right if c >= 5 else left
        r += 1
        for idx, li in enumerate(quote_data["line_items"], start=1):
            vals = [idx, str(li.get("description", "")), str(li.get("qty", "")),
                    str(li.get("unit", "")), str(li.get("unit_price", "")),
                    str(li.get("line_total", ""))]
            for c, v in enumerate(vals, start=1):
                cell = ws.cell(row=r, column=c, value=v)
                cell.font = dark
                cell.border = border
                cell.alignment = right if c >= 5 or c == 3 else (
                    wrap if c == 2 else left)
                if idx % 2 == 0:
                    cell.fill = light_fill
            r += 1
        r += 1

    # ---- 4. Travel ----
    if quote_data.get("travel_items"):
        _section("4. TRAVEL / LODGING / EQUIPMENT (ALLOWANCES)")
        for tv in quote_data["travel_items"]:
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=5)
            a = ws.cell(row=r, column=1, value=str(tv.get("description", "")))
            a.font = dark; a.alignment = left
            b = ws.cell(row=r, column=6, value=str(tv.get("amount", "")))
            b.font = dark; b.alignment = right
            r += 1
        r += 1

    # ---- 5. Total ----
    _section("5. TOTAL PRICE")

    def _tot(label, value, bold=False):
        nonlocal r
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=5)
        a = ws.cell(row=r, column=1, value=label)
        a.alignment = right
        a.font = Font(bold=bold, color=NAVY_HEX if bold else DARK_HEX,
                      size=11 if bold else 10)
        b = ws.cell(row=r, column=6, value=str(value))
        b.alignment = right
        b.font = Font(bold=bold, color=NAVY_HEX if bold else DARK_HEX,
                      size=11 if bold else 10)
        r += 1

    if quote_data.get("labor_subtotal"):
        _tot("Labor subtotal", quote_data["labor_subtotal"])
    if quote_data.get("materials_subtotal"):
        _tot("Materials subtotal (JRS-furnished)", quote_data["materials_subtotal"])
    if quote_data.get("travel_subtotal"):
        _tot("Travel / lodging / equipment", quote_data["travel_subtotal"])
    # Total final: fila a ancho completo, para que un total_text largo y
    # descriptivo no se monte sobre la etiqueta.
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=NCOLS)
    fcell = ws.cell(row=r, column=1,
                    value="TOTAL:   " + str(quote_data.get("total_text", "TBD")))
    fcell.font = Font(bold=True, color=NAVY_HEX, size=12)
    fcell.alignment = Alignment(horizontal="right", vertical="center")
    r += 2

    # ---- 6-9 listas ----
    def _lista(titulo, items):
        nonlocal r
        if not items:
            return
        _section(titulo)
        for it in items:
            _wrapline("•  " + str(it))
        r += 1

    _lista("6. ASSUMPTIONS", quote_data.get("assumptions"))
    _lista("7. EXCLUSIONS", quote_data.get("exclusions"))
    _lista("8. CLARIFICATIONS / OPEN ITEMS", quote_data.get("clarifications"))
    _lista("9. TERMS", quote_data.get("terms"))

    # ---- 10. Authorization ----
    _section("10. AUTHORIZATION")
    for label in ("Accepted by:", "Name / Title:", "Date:"):
        _wrapline(f"{label}   _________________________")
    r += 1

    # ---- Compliance note ----
    if quote_data.get("compliance_note"):
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=NCOLS)
        cell = ws.cell(row=r, column=1,
                       value="COMPLIANCE NOTE: " + str(quote_data["compliance_note"]))
        cell.font = muted; cell.alignment = wrap
        r += 1

    ws.sheet_view.showGridLines = False
    ws.freeze_panes = "A2"
    from openpyxl.worksheet.properties import PageSetupProperties
    ws.page_setup.orientation = "portrait"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr = PageSetupProperties(fitToPage=True)
    wb.save(output_path)
    return output_path


# ===================================================================
# Prueba de aislamiento:  python quote_files.py
# Genera un DOCX y un XLSX de muestra con datos reales (caso H&M con
# material, para ejercitar materials_subtotal y el basis mixto).
# ===================================================================
if __name__ == "__main__":
    ejemplo = {
        "quote_number": "JRS-2026-HM-NP-001",
        "date": "June 30, 2026",
        "prepared_for": "General Contractor — H&M NorthPark Remodel",
        "project_name": "H&M Remodel — NorthPark Center",
        "store_number": "TBD",
        "location": "NorthPark Center, Dallas, TX",
        "prepared_by": "Richard Bodington / JRS Retail Services",
        "phone": "832-361-6551",
        "basis": "Labor + JRS-Furnished Materials",
        "project_summary": (
            "Commercial retail remodel for H&M at NorthPark Center, Dallas TX. "
            "JRS provides labor for patch/paint, ceiling tile replacement, "
            "glue-down LVT installation, and final clean/punchlist.\n"
            "On this project JRS furnishes paint, ceiling tile + grid clips, and "
            "LVT adhesive (broken out separately). Flooring material (LVT) is "
            "supplied by the client; JRS installs only. Pricing assumes standard "
            "10-hour production shifts in an occupied mall/customer-facing retail "
            "environment."
        ),
        "scope_items": [
            "Patch and prep wall surfaces, then apply two-coat retail-grade paint to ~8,000 SF wall area (paint furnished by JRS).",
            "Remove and replace ceiling tile to ~5,000 SF, including grid clips (ceiling tile and grid clips furnished by JRS).",
            "Install glue-down LVT to ~5,000 SF, including floor prep and adhesive (LVT material by client; adhesive by JRS; JRS install only).",
            "Final clean and punchlist of all work areas to retail-ready condition.",
        ],
        "line_items": [
            {"description": "Patch + two-coat retail paint — labor", "qty": "8,000", "unit": "SF", "unit_price": "$1.85", "line_total": "$14,800"},
            {"description": "Ceiling tile replacement — labor", "qty": "5,000", "unit": "SF", "unit_price": "$2.10", "line_total": "$10,500"},
            {"description": "Glue-down LVT install (occupied retail) — labor", "qty": "5,000", "unit": "SF", "unit_price": "$4.75", "line_total": "$23,750"},
            {"description": "Final clean and punchlist — labor allowance", "qty": "1", "unit": "LS", "unit_price": "$3,500", "line_total": "$3,500"},
            {"description": "Paint — JRS-furnished material (two coats, 8,000 SF)", "qty": "1", "unit": "Allow", "unit_price": "$5,200", "line_total": "$5,200"},
            {"description": "Ceiling tile + grid clips — JRS-furnished material (5,000 SF)", "qty": "1", "unit": "Allow", "unit_price": "$6,500", "line_total": "$6,500"},
            {"description": "LVT adhesive — JRS-furnished material (5,000 SF)", "qty": "1", "unit": "Allow", "unit_price": "$2,200", "line_total": "$2,200"},
        ],
        "labor_subtotal": "$52,550",
        "materials_subtotal": "$13,900",
        "total_text": "$66,450 (Labor $52,550 + JRS-Furnished Materials $13,900)",
        "assumptions": [
            "Pricing based on standard 10-hour production shifts.",
            "Quantities approximate per scope provided (8,000 SF wall paint, 5,000 SF ceiling, 5,000 SF LVT); final price subject to confirmed measured quantities.",
            "JRS furnishes paint, ceiling tile + grid clips, and LVT adhesive, shown as separate line items.",
            "LVT flooring material supplied by client; JRS installs only.",
            "Pricing reflects occupied mall / customer-facing retail conditions.",
            "Debris haul-off / dumpster by GC unless added in writing.",
        ],
        "exclusions": [
            "LVT flooring material (by client).",
            "Permits, taxes (unless legally required), engineering, and design.",
            "Licensed electrical, plumbing, HVAC, fire alarm, sprinkler, and low-voltage work.",
            "Merchandise moving unless listed.",
            "After-hours mall / security fees unless added in writing.",
            "Hazardous materials — asbestos, mold, lead, environmental.",
        ],
        "clarifications": [
            "Confirm work hours: mall after-hours / overnight vs. daytime.",
            "Confirm exact measured wall SF, ceiling SF, and LVT SF prior to final price lock.",
            "Confirm whether a lift is required for ceiling work at this ceiling height.",
        ],
        "terms": [
            "Quote valid 15 days unless noted.",
            "Change orders required for any work outside the listed scope.",
            "Schedule subject to crew availability and site readiness.",
            "Payment terms: per subcontract agreement.",
        ],
    }

    docx_out = generar_docx_cotizacion(ejemplo, "JRS_Quote_sample.docx")
    xlsx_out = generar_xlsx_cotizacion(ejemplo, "JRS_Quote_sample.xlsx")
    print(f"DOCX generado: {docx_out}")
    print(f"XLSX generado: {xlsx_out}")
